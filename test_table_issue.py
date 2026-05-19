# -*- coding: utf-8 -*-
"""
Test script to reproduce and fix the black square issue in table header cells.
"""

import sys
import os
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from converter import MarkdownConverter
from config_manager import ConfigManager


def test_table_black_square():
    print("=" * 60)
    print("Testing table black square issue")
    print("=" * 60)
    
    config = ConfigManager()
    converter = MarkdownConverter(config)
    
    test_md = """# Table Test

## Simple Table

| Header1 | Header2 | Header3 |
|---------|---------|---------|
| Data1   | Data2   | Data3   |
| **Bold** | Normal | *Italic* |

## Test with empty first cell

|   | Name | Value |
|---|------|-------|
| 1 | Test | OK |
"""
    
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write(test_md)
        md_path = f.name
    
    output_path = os.path.join(tempfile.gettempdir(), 'test_table_black_square.docx')
    
    try:
        result = converter.convert(md_path)
        
        if result:
            doc = Document(result)
            
            print(f"\nDocument created: {result}")
            print(f"\nTables found: {len(doc.tables)}")
            
            for table_idx, table in enumerate(doc.tables):
                print(f"\n{'='*60}")
                print(f"Table {table_idx + 1}:")
                print(f"  Rows: {len(table.rows)}, Cols: {len(table.columns)}")
                
                for row_idx, row in enumerate(table.rows):
                    print(f"\n  Row {row_idx + 1}:")
                    for col_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        para_count = len(cell.paragraphs)
                        run_count = sum(len(p.runs) for p in cell.paragraphs)
                        
                        print(f"    Cell[{col_idx}]: '{cell_text[:30]}'")
                        print(f"      Paragraphs: {para_count}, Runs: {run_count}")
                        
                        if col_idx == 0 and row_idx == 0:
                            print(f"      [HEADER FIRST CELL]")
                            
                            # Check for potential issues
                            for p_idx, para in enumerate(cell.paragraphs):
                                print(f"      Para {p_idx}:")
                                print(f"        Text: '{para.text}'")
                                print(f"        Runs: {len(para.runs)}")
                                
                                for r_idx, run in enumerate(para.runs):
                                    run_text = repr(run.text)
                                    print(f"          Run {r_idx}: {run_text}")
                                    
                                    if run_text == "''" or run_text == '':
                                        print(f"            [WARNING] Empty run detected!")
            
            print(f"\n{'='*60}")
            print("Test completed!")
            print("Please open the document and check the first cell of each table header.")
            print(f"{'='*60}")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
    finally:
        if os.path.exists(md_path):
            os.unlink(md_path)


if __name__ == '__main__':
    test_table_black_square()
