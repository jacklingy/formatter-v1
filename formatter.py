"""
Word文档格式化模块 - WordFormatter类

本模块提供Word文档的专业格式化功能，支持：
- 标题格式化（6级标题样式自定义）
- 段落格式化（字体、行间距、首行缩进等）
- 列表格式化（有序/无序列表）
- 表格格式化（边框、背景色、对齐方式）
- 页码插入（使用Word域代码实现动态页码）
- 页面设置（页边距等）

主要依赖：
- python-docx：用于操作Word文档
- PyYAML：用于读取YAML配置文件

使用示例：
    from config_manager import ConfigManager
    formatter = WordFormatter(ConfigManager())
    output_path = formatter.format_document('input.docx')
"""

import os
import re
from docx import Document
from docx.shared import Pt, Twips, RGBColor, Emu, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class WordFormatter:
    """
    Word文档格式化器
    
    负责对Word文档进行全面的格式化处理，包括标题、段落、列表、表格和页码等。
    
    属性：
        config (ConfigManager): 配置管理器实例，用于读取格式配置
    
    使用流程：
        1. 实例化时传入ConfigManager对象
        2. 调用format_document()方法处理文档
        3. 返回格式化后的文档路径
    """
    
    def __init__(self, config_manager):
        """
        初始化WordFormatter实例
        
        参数：
            config_manager (ConfigManager): 配置管理器对象，
                用于读取format_config.yaml中的格式配置
        """
        self.config = config_manager
    
    def format_document(self, docx_file_path):
        """
        格式化Word文档的主入口方法
        
        执行完整的文档格式化流程：
        1. 验证输入文件是否存在
        2. 生成输出文件路径（在原文件名后添加"_已格式化"）
        3. 应用页面设置（页边距等）
        4. 格式化所有段落（标题、正文、列表）
        5. 格式化所有表格
        6. 插入页码
        7. 保存并返回输出路径
        
        参数：
            docx_file_path (str): 输入的Word文档路径
            
        返回：
            str: 格式化后的文档路径（xxx_已格式化.docx）
            
        异常：
            FileNotFoundError: 当输入文件不存在时抛出
            
        示例：
            >>> formatter = WordFormatter(config_manager)
            >>> output = formatter.format_document('report.docx')
            >>> print(output)
            'report_已格式化.docx'
        """
        if not os.path.exists(docx_file_path):
            raise FileNotFoundError(f"Word文档不存在: {docx_file_path}")

        output_path = self._get_output_path(docx_file_path)
        doc = Document(docx_file_path)

        self._apply_document_settings(doc)
        self._format_paragraphs(doc)
        self._format_tables(doc)
        self._insert_page_numbers(doc)

        doc.save(output_path)
        return output_path
    
    def _get_output_path(self, input_path):
        """
        生成输出文件路径
        
        在原始文件名的基础上添加"_已格式化"后缀，
        保持原有扩展名不变。
        
        参数：
            input_path (str): 输入文件路径
            
        返回：
            str: 输出文件路径，格式为"原文件名_已格式化.docx"
            
        示例：
            >>> self._get_output_path('/path/to/report.docx')
            '/path/to/report_已格式化.docx'
        """
        base_name = os.path.splitext(input_path)[0]
        return f"{base_name}_已格式化.docx"
    
    def _apply_document_settings(self, doc):
        """
        应用文档级别的页面设置
        
        从配置文件读取页边距设置，并应用到文档的所有节（section）。
        Word文档可以有多个节，每个节可以有不同的页面设置。
        
        支持的设置项：
        - top: 上边距（单位：twips，1英寸=1440twips）
        - bottom: 下边距
        - left: 左边距
        - right: 右边距
        
        默认值（如果配置中未指定）：
        - 上下边距：1440 twips (1英寸)
        - 左右边距：1800 twips (1.25英寸)
        
        参数：
            doc (Document): python-docx的Document对象
        """
        doc_settings = self.config.get_document_settings()
        if 'page_margin' in doc_settings:
            margins = doc_settings['page_margin']
            for section in doc.sections:
                section.top_margin = Twips(margins.get('top', 1440))
                section.bottom_margin = Twips(margins.get('bottom', 1440))
                section.left_margin = Twips(margins.get('left', 1800))
                section.right_margin = Twips(margins.get('right', 1800))
    
    def _format_paragraphs(self, doc):
        """
        格式化文档中的所有段落
        
        遍历文档中的每个段落，根据段落的类型应用不同的格式：
        - 标题段落 → 应用标题格式（调用_apply_heading_format）
        - 列表段落 → 应用列表格式（调用_apply_list_format）
        - 普通段落 → 应用正文格式（调用_apply_paragraph_format）
        
        段落类型判断规则：
        1. 检查样式名称是否包含"Heading"
        2. 检查是否为列表样式（List Paragraph）
        3. 检查XML属性是否包含列表标记
        4. 其余情况视为普通段落
        
        参数：
            doc (Document): python-docx的Document对象
        """
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            
            if self._is_heading_style(style_name):
                level = self._extract_heading_level(style_name)
                self._apply_heading_format(para, level)
            elif self._is_list_paragraph(para):
                self._apply_list_format(para)
            else:
                self._apply_paragraph_format(para)
    
    def _is_heading_style(self, style_name):
        """
        判断段落样式是否为标题样式
        
        通过正则表达式匹配样式名称，支持以下格式：
        - "Heading 1", "Heading 2", ... "Heading 6"
        - "heading 1", "heading 2" (大小写不敏感)
        - "Heading1", "Heading2" (空格可选)
        
        参数：
            style_name (str): 段落的样式名称
            
        返回：
            bool: 如果是标题样式返回True，否则返回False
        """
        return bool(re.match(r'Heading\s*\d+', style_name, re.IGNORECASE))
    
    def _extract_heading_level(self, style_name):
        """
        从样式名称中提取标题级别
        
        从"Heading N"格式的样式中提取数字N，
        表示这是第几级标题（1-6级）。
        
        参数：
            style_name (str): 标题样式名称，如"Heading 2"
            
        返回：
            int: 标题级别（1-6），如果无法提取则默认返回1
            
        示例：
            >>> self._extract_heading_level('Heading 3')
            3
            >>> self._extract_heading_level('Heading 1')
            1
        """
        match = re.search(r'\d+', style_name)
        return int(match.group()) if match else 1
    
    def _is_list_paragraph(self, para):
        """
        判断段落是否为列表段落
        
        综合多种方式判断一个段落是否属于列表：
        
        1. **样式名称检查**：样式名包含"List"关键字
        2. **XML属性检查**：
           - pPr.numPr 存在（编号列表属性）
           - pPr.bulletListPr 存在（项目符号属性）
        3. **文本内容检查**：
           - 以"•"开头（项目符号）
           - 以"-"开头（短横线项目符号）
           - 以数字+点/括号开头（有序列表，如"1."、"2)"）
        
        这种多重检查机制确保了即使在不同版本的Word中创建的列表，
        都能被正确识别。
        
        参数：
            para (Paragraph): python-docx的Paragraph对象
            
        返回：
            bool: 如果是列表段落返回True，否则返回False
        """
        if para.style and 'List' in para.style.name:
            return True
        if para._element.pPr is not None:
            pPr = para._element.pPr
            if pPr.numPr is not None:
                return True
            try:
                if hasattr(pPr, 'bulletListPr') and pPr.bulletListPr is not None:
                    return True
            except AttributeError:
                pass
        text = para.text.strip()
        if text and (text.startswith('•') or text.startswith('-') or 
                     re.match(r'^\d+[.\)]', text)):
            return True
        return False
    
    def _apply_heading_format(self, para, level):
        """
        应用标题格式到段落
        
        根据配置文件中对应级别的标题样式设置，格式化标题段落。
        支持6级标题（level 1-6），每级可独立配置。
        
        可配置的格式属性：
        - font_name: 字体名称（默认：黑体）
        - font_size: 字号（单位：磅，默认：14）
        - bold: 是否加粗（默认：True）
        - alignment: 对齐方式（left/center/right/justify）
        - space_before: 段前间距（单位：twips，默认：120）
        - space_after: 段后间距（单位：twips，默认：60）
        - color: 字体颜色（十六进制，默认：000000黑色）
        - first_line_indent: 首行缩进（单位：twips，默认：480即2字符）
        
        技术细节：
        - 使用rFonts同时设置中西文字体，确保中文显示正确
        - 需要遍历paragraph的所有run并分别设置字体
        
        参数：
            para (Paragraph): 要格式化的段落对象
            level (int): 标题级别（1-6）
        """
        style_config = self.config.get_heading_style(level)
        
        font_name = style_config.get('font_name', '黑体')
        font_size = style_config.get('font_size', 14)
        bold = style_config.get('bold', True)
        alignment = style_config.get('alignment', 'left')
        space_before = style_config.get('space_before', 120)
        space_after = style_config.get('space_after', 60)
        color = style_config.get('color', '000000')
        first_line_indent = style_config.get('first_line_indent', 480)

        for run in para.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)

        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        para.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)

        para_format = para.paragraph_format
        para_format.space_before = Pt(space_before / 20)
        para_format.space_after = Pt(space_after / 20)
        
        if first_line_indent and first_line_indent > 0:
            para_format.first_line_indent = Twips(first_line_indent)
        else:
            para_format.first_line_indent = Twips(0)
    
    def _apply_paragraph_format(self, para):
        """
        应用普通段落格式
        
        为非标题、非列表的普通正文段落应用标准格式。
        这是文档中最常用的格式，决定了正文的整体外观。
        
        可配置的格式属性：
        - default_font: 正文字体（默认：宋体）
        - default_size: 正文字号（单位：磅，默认：12即小四）
        - default_color: 字体颜色（默认：000000黑色）
        - line_spacing: 行距（倍数，默认：1.5倍行距）
        - first_line_indent: 首行缩进（单位：twips，默认：480即2字符）
        - alignment: 对齐方式（默认：左对齐，通常设置为两端对齐justify）
        - bold: 是否加粗（默认：False）
        - space_before: 段前间距（单位：磅，默认：0）
        - space_after: 段后间距（单位：磅，默认：0）
        
        中文排版规范建议：
        - 首行缩进2字符（约480 twips）
        - 行距1.5倍或固定值28磅
        - 两端对齐（justify）使右边缘整齐
        
        参数：
            para (Paragraph): 要格式化的普通段落对象
        """
        para_style = self.config.get_paragraph_style()
        font_settings = self.config.get_font_settings()

        font_name = font_settings.get('default_font', '宋体')
        font_size = font_settings.get('default_size', 12)
        color = font_settings.get('default_color', '000000')
        line_spacing = para_style.get('line_spacing', 1.5)
        first_indent = para_style.get('first_line_indent', 480)
        alignment = para_style.get('alignment', 'left')
        bold = para_style.get('bold', False)
        space_before = para_style.get('space_before', 0)
        space_after = para_style.get('space_after', 0)

        for run in para.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(color)
            run.font.bold = bold

        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        para_format = para.paragraph_format
        para_format.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        para_format.line_spacing = line_spacing
        para_format.first_line_indent = Twips(first_indent)
        para_format.space_before = Pt(space_before / 20)
        para_format.space_after = Pt(space_after / 20)
    
    def _apply_list_format(self, para):
        """
        应用列表段落格式
        
        为有序列表和无序列表项应用专门的格式。
        列表项通常需要左侧缩进以体现层级关系。
        
        可配置的格式属性：
        - 字体和字号：继承自全局字体设置
        - indent_left: 左侧缩进（单位：twips，默认：720即3字符）
          - 用于控制列表项的缩进深度
          - 值越大，列表项越靠右显示
        
        注意事项：
        - 列表的编号/符号由Word自动管理，此处只设置文本格式
        - 不设置首行缩进（与正文不同）
        
        参数：
            para (Paragraph): 要格式化的列表段落对象
        """
        list_style = self.config.get_list_style()
        font_settings = self.config.get_font_settings()

        font_name = font_settings.get('default_font', '宋体')
        font_size = font_settings.get('default_size', 12)
        indent_left = list_style.get('indent_left', 720)

        for run in para.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)

        para.paragraph_format.left_indent = Twips(indent_left)
    
    def _format_tables(self, doc):
        """
        格式化文档中的所有表格
        
        遍历文档中的每个表格，统一应用表格样式。
        表格格式的重点是提升专业性和可读性。
        
        格式化内容包括：
        1. 表格居中对齐
        2. 自动调整列宽（自适应内容）
        3. 区分表头和数据行的样式
        4. 设置单元格边框（实线黑色细边框）
        5. 设置单元格背景色（表头可不同）
        6. 单元格内文字垂直居中
        7. 统一字体和字号
        
        配置结构：
        - header: 表头行样式（第一行）
          - 通常加粗、可能有背景色
        - body: 数据行样式（其余行）
          - 通常不加粗、白色背景
        
        技术要点：
        - 需要直接操作底层XML元素来设置边框和背景
        - 使用OxmlElement创建Word XML元素
        
        参数：
            doc (Document): python-docx的Document对象
        """
        table_style_config = self.config.get_table_style()
        
        for table in doc.tables:
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            self._set_table_auto_fit(table)
            
            for row_idx, row in enumerate(table.rows):
                is_header_row = (row_idx == 0)
                
                if is_header_row:
                    style = table_style_config.get('header', {})
                else:
                    style = table_style_config.get('body', {})
                
                font_name = style.get('font_name', '宋体')
                font_size = style.get('font_size', 10.5)
                bold = style.get('bold', False)
                bg_color = style.get('background_color', [255, 255, 255])
                text_align = style.get('text_alignment', 'center')
                
                row.height_rule = None
                
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    self._set_cell_borders(cell)
                    
                    self._set_cell_background(cell, bg_color)
                    
                    for paragraph in cell.paragraphs:
                        if is_header_row:
                            paragraph.paragraph_format.keep_with_next = True
                        
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                            run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
                            run._element.rPr.rFonts.set(qn('w:cs'), font_name)
                            
                            run.font.size = Pt(font_size)
                            run.font.bold = bold
                            run.font.italic = False
                            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def _set_table_auto_fit(self, table):
        """
        设置表格自动适应列宽
        
        在表格的XML属性中添加tblLayout元素，
        设置type="auto"使表格列宽根据内容自动调整。
        这比固定列宽更灵活，能更好地适应不同长度的文本。
        
        实现原理：
        - 操作table底层的_tbl元素（XML表示）
        - 创建w:tblLayout元素并设置属性
        - 如果tblPr不存在则创建它
        
        参数：
            table (Table): python-docx的Table对象
        """
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'auto')
        tblPr.append(tblLayout)
        
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def _set_cell_borders(self, cell, border_color='000000', border_width=1):
        """
        设置单元格边框
        
        为表格单元格添加四周边框（上、下、左、右）。
        使用Word的tcBorders XML元素实现。
        
        边框属性：
        - val: 边框样式（single=单实线）
        - sz: 边框宽度（单位：1/8磅，border_width * 8）
          - 默认值1对应1/8磅（极细线）
          - 常用值4对应0.5磅（中等粗细）
        - color: 边框颜色（十六进制，默认黑色）
        - space: 边框与内容的间距（默认0）
        
        实现步骤：
        1. 获取或创建单元格的tcPr（表格单元格属性）元素
        2. 创建tcBorders元素
        3. 为四个方向分别创建边框元素并设置属性
        4. 将边框元素添加到tcPr中
        
        参数：
            cell (TableCell): python-docx的TableCell对象
            border_color (str): 边框颜色，十六进制格式（默认'000000'黑色）
            border_width (int): 边框宽度系数（默认1）
        """
        tc = cell._tc
        tcPr = tc.tcPr if tc.tcPr is not None else OxmlElement('w:tcPr')
        
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(border_width * 8))
            border.set(qn('w:color'), border_color)
            border.set(qn('w:space'), '0')
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
        
        if tc.tcPr is None:
            tc.insert(0, tcPr)
    
    def _set_cell_background(self, cell, rgb_color):
        """
        设置单元格背景色
        
        为表格单元格设置填充色（背景色）。
        常用于区分表头和数据行。
        
        参数说明：
        - rgb_color (list/tuple): RGB颜色值，格式为[R, G, B]
          - 每个分量范围：0-255
          - 例如：[240, 240, 240] 浅灰色
          - 例如：[255, 255, 255] 白色
          - 例如：[200, 220, 240] 浅蓝色
        
        实现原理：
        - 使用w:shd（shading）元素设置底纹
        - val='clear'表示清除样式
        - color='auto'表示前景色自动
        - fill设置实际的背景色（十六进制RGB值）
        
        参数：
            cell (TableCell): python-docx的TableCell对象
            rgb_color (list): RGB颜色值列表 [R, G, B]，每个值0-255
        """
        tc = cell._tc
        tcPr = tc.tcPr if tc.tcPr is not None else OxmlElement('w:tcPr')
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        
        r, g, b = rgb_color
        hex_color = '{:02X}{:02X}{:02X}'.format(r, g, b)
        shd.set(qn('w:fill'), hex_color)
        
        tcPr.append(shd)
        
        if tc.tcPr is None:
            tc.insert(0, tcPr)

    def _insert_page_numbers(self, doc):
        """
        插入动态页码到文档页脚
        
        在文档每个节的页脚位置插入页码字段。
        使用Word域代码（Field Code）实现真正的动态页码，
        而不是静态数字。这样当文档编辑后页码会自动更新。
        
        功能特点：
        - 动态页码：使用PAGE域代码，会随文档变化自动更新
        - 可配置位置：目前支持底部居中（bottom_center）
        - 可配置字体：支持分别设置中文和西文字体
        - 多节支持：为每个节独立设置页脚
        - 首页显示选项：可选择是否在首页显示页码
        
        配置项（从config读取）：
        - enabled: 是否启用页码功能（必须为True才插入）
        - position: 页码位置（bottom_center=底部居中）
        - font_name: 主字体名称
        - chinese_font: 中文字体（用于东亚字符）
        - western_font: 西文字体（用于数字和英文）
        - font_size: 字号（单位：磅，默认10.5即五号字）
        - bold: 是否加粗
        - format: 页码格式（如'{n}'、'第{n}页'等）
        - start_from: 起始页码（默认1）
        - show_on_first_page: 首页是否显示（默认True）
        
        技术实现：
        - 使用fldChar元素标记域的开始、分隔和结束
        - instrText包含实际的域代码（PAGE）
        - 域代码会被Word解析并显示为实际页码
        
        参数：
            doc (Document): python-docx的Document对象
        """
        page_number_config = self.config.get_page_number_settings()
        
        if not page_number_config.get('enabled', False):
            return
        
        position = page_number_config.get('position', 'bottom_center')
        font_name = page_number_config.get('font_name', '宋体')
        chinese_font = page_number_config.get('chinese_font', '宋体')
        western_font = page_number_config.get('western_font', 'Times New Roman')
        font_size = page_number_config.get('font_size', 10.5)
        bold = page_number_config.get('bold', False)
        format_str = page_number_config.get('format', '{n}')
        start_from = page_number_config.get('start_from', 1)
        show_on_first = page_number_config.get('show_on_first_page', True)
        
        for section_idx, section in enumerate(doc.sections):
            footer = section.footer
            footer.is_linked_to_previous = False
            
            if len(footer.paragraphs) == 0:
                footer.add_paragraph()
            
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self._add_page_number_field(paragraph, start_from, font_name, chinese_font, 
                                       western_font, font_size, bold, format_str)
    
    def _add_page_number_field(self, paragraph, start_from, font_name, chinese_font, 
                               western_font, font_size, bold, format_str):
        """
        在段落中添加页码域（Field）
        
        构建Word域代码结构的完整XML元素序列。
        Word域由以下部分组成：
        
        1. **fldChar begin**：域开始标记
           - 告诉Word这里开始一个域
           
        2. **instrText**：域指令文本
           - 包含实际的域代码（如" PAGE "）
           - PAGE是Word内置域，用于显示当前页码
           - 可选参数\\* MERGEFORMAT保持格式
           
        3. **fldChar separate**：域分隔符
           - 分隔域代码和域结果显示区域
           - 分隔符之前是代码，之后是显示的内容
           
        4. **fldChar end**：域结束标记
           - 标记域的结束
        
        字体设置技巧：
        - 分别设置eastAsia（中文）、ascii/hAnsi/cs（西文）字体
        - 确保页码数字使用西文字体（如Times New Roman）
        - 确保页码前后如果有中文则使用中文字体
        
        参数：
            paragraph (Paragraph): 要添加页码的段落（通常是页脚段落）
            start_from (int): 起始页码（预留参数，暂未完全实现）
            font_name (str): 主字体名称
            chinese_font (str): 中文字体名称
            western_font (str): 西文字体名称（用于数字显示）
            font_size (int): 字号（单位：磅）
            bold (bool): 是否加粗
            format_str (str): 页码格式模板（预留参数）
        """
        run = paragraph.add_run()
        
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        field_code = ' PAGE '
        if start_from != 1:
            field_code = f' PAGE \\* MERGEFORMAT '
        instr_text.text = field_code
        
        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')
        
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)
        
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
        run._element.rPr.rFonts.set(qn('w:ascii'), western_font)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), western_font)
        run._element.rPr.rFonts.set(qn('w:cs'), western_font)
        
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
