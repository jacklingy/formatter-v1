import os
import re
import urllib.request
from docx import Document
from docx.shared import Pt, Twips, RGBColor, Emu, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

class MarkdownConverter:
    def __init__(self, config_manager):
        self.config = config_manager

    def convert(self, md_file_path):
        if not os.path.exists(md_file_path):
            raise FileNotFoundError(f"Markdown文件不存在: {md_file_path}")

        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        output_path = self._get_output_path(md_file_path)
        doc = Document()
        
        self._apply_document_settings(doc)
        self._parse_markdown(doc, md_content)
        self._insert_page_numbers(doc)
        
        doc.save(output_path)
        return output_path

    def _get_output_path(self, input_path):
        base_name = os.path.splitext(input_path)[0]
        return f"{base_name}_格式化.docx"

    def _apply_document_settings(self, doc):
        doc_settings = self.config.get_document_settings()
        if 'page_margin' in doc_settings:
            margins = doc_settings['page_margin']
            for section in doc.sections:
                section.top_margin = Twips(margins.get('top', 1440))
                section.bottom_margin = Twips(margins.get('bottom', 1440))
                section.left_margin = Twips(margins.get('left', 1800))
                section.right_margin = Twips(margins.get('right', 1800))

    def _parse_markdown(self, doc, content):
        lines = content.split('\n')
        i = 0
        in_list = False
        list_type = None
        list_items = []
        in_blockquote = False
        blockquote_lines = []

        while i < len(lines):
            line = lines[i].rstrip()

            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                if in_list and list_items:
                    self._add_list(doc, list_type, list_items)
                    list_items = []
                    in_list = False
                if in_blockquote and blockquote_lines:
                    self._add_blockquote(doc, '\n'.join(blockquote_lines))
                    blockquote_lines = []
                    in_blockquote = False
                
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                self._add_heading(doc, text, level)
                i += 1
                continue

            image_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', line)
            if image_match:
                if in_list and list_items:
                    self._add_list(doc, list_type, list_items)
                    list_items = []
                    in_list = False
                if in_blockquote and blockquote_lines:
                    self._add_blockquote(doc, '\n'.join(blockquote_lines))
                    blockquote_lines = []
                    in_blockquote = False
                
                alt_text = image_match.group(1)
                image_url = image_match.group(2)
                self._add_image(doc, image_url, alt_text)
                i += 1
                continue

            if line.startswith('---') or line.startswith('***') or line.startswith('___'):
                if in_list and list_items:
                    self._add_list(doc, list_type, list_items)
                    list_items = []
                    in_list = False
                if in_blockquote and blockquote_lines:
                    self._add_blockquote(doc, '\n'.join(blockquote_lines))
                    blockquote_lines = []
                    in_blockquote = False
                    
                self._add_horizontal_rule(doc)
                i += 1
                continue

            if line.startswith('>'):
                if in_list and list_items:
                    self._add_list(doc, list_type, list_items)
                    list_items = []
                    in_list = False
                
                if not in_blockquote:
                    in_blockquote = True
                    blockquote_lines = []
                
                quote_text = line[1:].strip() if len(line) > 1 else ''
                blockquote_lines.append(quote_text)
                i += 1
                continue

            task_match = re.match(r'^- \[([ xX])\]\s*(.+)$', line)
            if task_match:
                if in_blockquote and blockquote_lines:
                    self._add_blockquote(doc, '\n'.join(blockquote_lines))
                    blockquote_lines = []
                    in_blockquote = False
                    
                if in_list and list_type != 'task':
                    self._add_list(doc, list_type, list_items)
                    list_items = []
                
                in_list = True
                list_type = 'task'
                checked = task_match.group(1).strip().lower() == 'x'
                item_text = task_match.group(2).strip()
                list_items.append({'text': item_text, 'checked': checked})
                i += 1
                continue

            stripped_line = line.lstrip()
            
            if stripped_line.startswith('- ') or stripped_line.startswith('* '):
                if in_blockquote and blockquote_lines:
                    self._add_blockquote(doc, '\n'.join(blockquote_lines))
                    blockquote_lines = []
                    in_blockquote = False
                    
                if not in_list or list_type != 'bullet':
                    if in_list and list_items:
                        self._add_list(doc, list_type, list_items)
                        list_items = []
                    in_list = True
                    list_type = 'bullet'
                item_text = stripped_line[2:].strip()
                indent_level = len(line) - len(stripped_line)
                list_items.append({'text': item_text, 'indent': indent_level})
                i += 1
                continue

            numbered_match = re.match(r'^\d+\.\s+(.+)$', line)
            if numbered_match:
                if in_blockquote and blockquote_lines:
                    self._add_blockquote(doc, '\n'.join(blockquote_lines))
                    blockquote_lines = []
                    in_blockquote = False
                    
                if not in_list or list_type != 'numbered':
                    if in_list and list_items:
                        self._add_list(doc, list_type, list_items)
                        list_items = []
                    in_list = True
                    list_type = 'numbered'
                list_items.append(numbered_match.group(1).strip())
                i += 1
                continue

            if in_list and list_items:
                self._add_list(doc, list_type, list_items)
                list_items = []
                in_list = False

            if in_blockquote and blockquote_lines:
                self._add_blockquote(doc, '\n'.join(blockquote_lines))
                blockquote_lines = []
                in_blockquote = False

            if line.strip() == '':
                i += 1
                continue

            code_block_match = re.match(r'^```(\w*)$', line)
            if code_block_match:
                lang = code_block_match.group(1)
                code_lines = []
                i += 1
                while i < len(lines) and lines[i].strip() != '```':
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block(doc, '\n'.join(code_lines), lang)
                i += 1
                continue

            table_match = re.match(r'^\|(.+)\|$', line)
            if table_match:
                table_lines = [line]
                i += 1
                while i < len(lines) and (lines[i].startswith('|') or re.match(r'^[\|\s\-:]+$', lines[i])):
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(doc, table_lines)
                continue

            text = line
            while i + 1 < len(lines) and lines[i + 1].strip() != '' and \
                  not lines[i + 1].startswith('#') and \
                  not lines[i + 1].startswith('-') and \
                  not lines[i + 1].startswith('*') and \
                  not re.match(r'^\d+\.', lines[i + 1]) and \
                  not lines[i + 1].startswith('|') and \
                  not lines[i + 1].startswith('```') and \
                  not lines[i + 1].startswith('>') and \
                  not (lines[i + 1].startswith('---') or lines[i + 1].startswith('***')) and \
                  not re.match(r'^!\[.*\]\(.+\)', lines[i + 1]):
                i += 1
                text += '\n' + lines[i]

            self._add_formatted_paragraph(doc, text)
            i += 1

        if in_list and list_items:
            self._add_list(doc, list_type, list_items)
        if in_blockquote and blockquote_lines:
            self._add_blockquote(doc, '\n'.join(blockquote_lines))

    def _add_heading(self, doc, text, level):
        style_config = self.config.get_heading_style(level)
        
        para = doc.add_paragraph()
        self._apply_inline_formatting(para, text)

        font_name = style_config.get('font_name', '黑体')
        font_size = style_config.get('font_size', 14)
        bold = style_config.get('bold', True)
        space_before = style_config.get('space_before', 120)
        space_after = style_config.get('space_after', 60)
        color = style_config.get('color', '000000')
        first_line_indent = style_config.get('first_line_indent', 480)

        for run in para.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)

        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        para_format = para.paragraph_format
        para_format.space_before = Pt(space_before / 20)
        para_format.space_after = Pt(space_after / 20)
        
        if first_line_indent and first_line_indent > 0:
            para_format.first_line_indent = Twips(first_line_indent)
        else:
            para_format.first_line_indent = Twips(0)

    def _add_formatted_paragraph(self, doc, text):
        para_style = self.config.get_paragraph_style()
        font_settings = self.config.get_font_settings()

        para = doc.add_paragraph()
        self._apply_inline_formatting(para, text)

        font_name = font_settings.get('default_font', '宋体')
        font_size = font_settings.get('default_size', 12)
        color = font_settings.get('default_color', '000000')
        line_spacing = para_style.get('line_spacing', 1.5)
        first_indent = para_style.get('first_line_indent', 480)
        space_before = para_style.get('space_before', 0)
        space_after = para_style.get('space_after', 0)

        for run in para.runs:
            if not run.font.name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if not run.font.size:
                run.font.size = Pt(font_size)
            if not run.font.color.rgb or str(run.font.color.rgb) == '000000':
                run.font.color.rgb = RGBColor.from_string(color)

        para_format = para.paragraph_format
        para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_format.line_spacing = line_spacing
        para_format.first_line_indent = Twips(first_indent)
        para_format.space_before = Pt(space_before / 20)
        para_format.space_after = Pt(space_after / 20)

    def _apply_inline_formatting(self, paragraph, text):
        pattern = r'(\*\*.*?\*\*|\*.*?\*|~~.*?~~|`[^`]+|!\[([^\]]*)\]\(([^)]+)\)|\[([^\]]+)\]\(([^)]+)\))'
        parts = re.split(pattern, text)
        
        i = 0
        while i < len(parts):
            part = parts[i]
            
            if not part:
                i += 1
                continue
            
            if part.startswith('**') and part.endswith('**'):
                run_text = part[2:-2]
                run = paragraph.add_run(run_text)
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run_text = part[1:-1]
                run = paragraph.add_run(run_text)
                run.italic = True
            elif part.startswith('~~') and part.endswith('~~'):
                run_text = part[2:-2]
                run = paragraph.add_run(run_text)
                run.font.strike = True
            elif part.startswith('`') and part.endswith('`'):
                run_text = part[1:-1]
                run = paragraph.add_run(run_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            elif part.startswith('![') and '](http' in part or '](./' in part:
                img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', part)
                if img_match:
                    alt_text = img_match.group(1) or ''
                    url = img_match.group(2) or ''
                    self._add_inline_image(paragraph, url, alt_text)
                else:
                    run = paragraph.add_run(part)
            elif part.startswith('[') and '](http' in part or '](./' in part:
                link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
                if link_match:
                    link_text = link_match.group(1) or ''
                    url = link_match.group(2) or ''
                    self._add_hyperlink_to_paragraph(paragraph, link_text, url)
                else:
                    run = paragraph.add_run(part)
            else:
                run = paragraph.add_run(part)
            
            i += 1

    def _add_image(self, doc, image_url, alt_text=''):
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if image_url.startswith(('http://', 'https://')):
                temp_path = f'temp_image_{hash(image_url)}.png'
                try:
                    urllib.request.urlretrieve(image_url, temp_path)
                    run = para.add_run()
                    run.add_picture(temp_path, width=Inches(4.5))
                    
                    if alt_text:
                        caption_para = doc.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_para.add_run(alt_text)
                        caption_run.font.size = Pt(9)
                        caption_run.font.color.rgb = RGBColor(100, 100, 100)
                        
                finally:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
            elif os.path.exists(image_url):
                run = para.add_run()
                run.add_picture(image_url, width=Inches(4.5))
                
                if alt_text:
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(alt_text)
                    caption_run.font.size = Pt(9)
                    caption_run.font.color.rgb = RGBColor(100, 100, 100)
            else:
                error_para = doc.add_paragraph()
                error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                error_run = error_para.add_run(f'[图片未找到: {alt_text or image_url}]')
                error_run.font.color.rgb = RGBColor(255, 0, 0)
                error_run.font.italic = True
                
        except Exception as e:
            error_para = doc.add_paragraph()
            error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            error_run = error_para.add_run(f'[图片加载失败: {str(e)}]')
            error_run.font.color.rgb = RGBColor(255, 0, 0)
            error_run.font.italic = True

    def _add_inline_image(self, paragraph, image_url, alt_text=''):
        try:
            run = paragraph.add_run()
            
            if image_url.startswith(('http://', 'https://')):
                temp_path = f'temp_inline_image_{hash(image_url)}.png'
                try:
                    urllib.request.urlretrieve(image_url, temp_path)
                    run.add_picture(temp_path, height=Pt(16))
                finally:
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
            elif os.path.exists(image_url):
                run.add_picture(image_url, height=Pt(16))
            else:
                text_run = paragraph.add_run(f'[{alt_text or "图片"}]')
                text_run.font.color.rgb = RGBColor(255, 0, 0)
                return
                
        except Exception as e:
            text_run = paragraph.add_run(f'[{alt_text or "图片"}]')
            text_run.font.color.rgb = RGBColor(255, 0, 0)

    def _add_hyperlink_to_paragraph(self, paragraph, text, url):
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), url)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)

        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        rel = paragraph.part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    def _add_list(self, doc, list_type, items):
        list_style = self.config.get_list_style()
        font_settings = self.config.get_font_settings()

        if list_type == 'task':
            for item in items:
                para = doc.add_paragraph()
                
                checkbox = '☑' if item['checked'] else '☐'
                self._apply_inline_formatting(para, f"{checkbox} {item['text']}")

                font_name = font_settings.get('default_font', '宋体')
                font_size = font_settings.get('default_size', 12)

                for run in para.runs:
                    if not run.font.name or run.font.name == 'Calibri':
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    if not run.font.size:
                        run.font.size = Pt(font_size)

                indent_left = list_style.get('indent_left', 720)
                para.paragraph_format.left_indent = Twips(indent_left)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            for item in items:
                para = doc.add_paragraph()
                
                if isinstance(item, dict):
                    item_text = item['text']
                    indent_level = item.get('indent', 0)
                else:
                    item_text = item
                    indent_level = 0
                
                if list_type == 'bullet':
                    prefix = list_style.get('bullet_char', '•') + ' '
                else:
                    prefix = ''
                    
                self._apply_inline_formatting(para, f"{prefix}{item_text}")

                font_name = font_settings.get('default_font', '宋体')
                font_size = font_settings.get('default_size', 12)

                for run in para.runs:
                    if not run.font.name or run.font.name == 'Calibri':
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    if not run.font.size:
                        run.font.size = Pt(font_size)

                base_indent = list_style.get('indent_left', 720)
                total_indent = base_indent + (indent_level * 360)
                para.paragraph_format.left_indent = Twips(total_indent)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_blockquote(self, doc, text):
        para = doc.add_paragraph()
        self._apply_inline_formatting(para, text)

        font_settings = self.config.get_font_settings()
        font_name = font_settings.get('default_font', '宋体')
        font_size = font_settings.get('default_size', 11)

        for run in para.runs:
            if not run.font.name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if not run.font.size:
                run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(80, 80, 80)
            run.italic = True

        para.paragraph_format.left_indent = Twips(720)
        para.paragraph_format.right_indent = Twips(720)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        border_para = OxmlElement('w:pBdr')
        left_border = OxmlElement('w:left')
        left_border.set(qn('w:val'), 'single')
        left_border.set(qn('w:sz'), '18')
        left_border.set(qn('w:color'), '808080')
        left_border.set(qn('w:space'), '6')
        border_para.append(left_border)
        para._p.get_or_add_pPr().append(border_para)

    def _add_horizontal_rule(self, doc):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)

        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), 'A0A0A0')
        bottom.set(qn('w:space'), '2')
        pBdr.append(bottom)
        para._p.get_or_add_pPr().append(pBdr)

    def _add_code_block(self, doc, code, language=None):
        para = doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(40, 40, 40)
        para.paragraph_format.left_indent = Twips(720)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5')
        para._p.get_or_add_pPr().append(shd)

    def _add_table(self, doc, table_lines):
        if not table_lines or len(table_lines) < 1:
            return

        rows_data = []
        col_alignments = []
        
        separator_idx = None
        for idx, line in enumerate(table_lines):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            if re.match(r'^[\|\s\-:]+$', line):
                separator_idx = idx
                for cell in cells:
                    cell = cell.strip()
                    if ':' in cell:
                        if cell.startswith(':') and cell.endswith(':'):
                            col_alignments.append('center')
                        elif cell.endswith(':'):
                            col_alignments.append('right')
                        else:
                            col_alignments.append('left')
                    else:
                        col_alignments.append('center')
                continue
            
            rows_data.append(cells)

        if not rows_data:
            return

        num_cols = len(rows_data[0])
        
        while len(col_alignments) < num_cols:
            col_alignments.append('center')

        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        self._set_table_auto_fit(table)
        
        table_style_config = self.config.get_table_style()

        for row_idx, row_data in enumerate(rows_data):
            row = table.rows[row_idx]
            is_header_row = (row_idx == 0)
            
            if is_header_row:
                style = table_style_config.get('header', {})
            else:
                style = table_style_config.get('body', {})
            
            font_name = style.get('font_name', '宋体')
            font_size = style.get('font_size', 10.5)
            bold = style.get('bold', False)
            bg_color = style.get('background_color', [255, 255, 255])
            
            row.height_rule = None
            
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= len(row.cells):
                    continue
                    
                cell = row.cells[col_idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                self._set_cell_borders(cell)
                self._set_cell_background(cell, bg_color)
                
                cell.text = ''
                
                para = cell.paragraphs[0]
                
                for run in para.runs:
                    run._element.getparent().remove(run._element)
                
                self._apply_inline_formatting(para, cell_text)
                
                col_alignment = col_alignments[col_idx] if col_idx < len(col_alignments) else 'center'
                
                align_map = {
                    'left': WD_ALIGN_PARAGRAPH.LEFT,
                    'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT
                }
                cell_alignment = align_map.get(col_alignment, WD_ALIGN_PARAGRAPH.CENTER)
                
                for paragraph in cell.paragraphs:
                    if is_header_row:
                        paragraph.paragraph_format.keep_with_next = True
                    
                    paragraph.alignment = cell_alignment
                    
                    for run in paragraph.runs:
                        if not run.font.name:
                            run.font.name = font_name
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                            run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
                            run._element.rPr.rFonts.set(qn('w:cs'), font_name)
                        
                        if not run.font.size:
                            run.font.size = Pt(font_size)
                        
                        if is_header_row and not run.bold:
                            run.font.bold = bold
                            
                        if not run.font.color.rgb or str(run.font.color.rgb) == '000000':
                            run.font.color.rgb = RGBColor(0, 0, 0)

    def _set_table_auto_fit(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'auto')
        tblPr.append(tblLayout)
        
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
    
    def _set_cell_borders(self, cell, border_color='000000', border_width=1):
        tc = cell._tc
        tcPr = tc.tcPr if tc.tcPr is not None else OxmlElement('w:tcPr')
        
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(border_width * 8))
            border.set(qn('w:color'), border_color)
            border.set(qn('w:space'), '0')
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
        
        if tc.tcPr is None:
            tc.insert(0, tcPr)
    
    def _set_cell_background(self, cell, rgb_color):
        tc = cell._tc
        tcPr = tc.tcPr if tc.tcPr is not None else OxmlElement('w:tcPr')
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        
        r, g, b = rgb_color
        hex_color = '{:02X}{:02X}{:02X}'.format(r, g, b)
        shd.set(qn('w:fill'), hex_color)
        
        tcPr.append(shd)
        
        if tc.tcPr is None:
            tc.insert(0, tcPr)

    def _insert_page_numbers(self, doc):
        page_number_config = self.config.get_page_number_settings()
        
        if not page_number_config.get('enabled', False):
            return
        
        position = page_number_config.get('position', 'bottom_center')
        font_name = page_number_config.get('font_name', '宋体')
        chinese_font = page_number_config.get('chinese_font', '宋体')
        western_font = page_number_config.get('western_font', 'Times New Roman')
        font_size = page_number_config.get('font_size', 10.5)
        bold = page_number_config.get('bold', False)
        format_str = page_number_config.get('format', '{n}')
        start_from = page_number_config.get('start_from', 1)
        show_on_first = page_number_config.get('show_on_first_page', True)
        
        for section_idx, section in enumerate(doc.sections):
            footer = section.footer
            footer.is_linked_to_previous = False
            
            if len(footer.paragraphs) == 0:
                footer.add_paragraph()
            
            paragraph = footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self._add_page_number_field(paragraph, start_from, font_name, chinese_font, 
                                       western_font, font_size, bold, format_str)

    def _add_page_number_field(self, paragraph, start_from, font_name, chinese_font, 
                               western_font, font_size, bold, format_str):
        run = paragraph.add_run()
        
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        field_code = ' PAGE '
        if start_from != 1:
            field_code = f' PAGE \\* MERGEFORMAT '
        instr_text.text = field_code
        
        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')
        
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)
        
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
        run._element.rPr.rFonts.set(qn('w:ascii'), western_font)
        run._element.rPr.rFonts.set(qn('w:hAnsi'), western_font)
        run._element.rPr.rFonts.set(qn('w:cs'), western_font)
        
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)
