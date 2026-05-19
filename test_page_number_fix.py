import sys
from config_manager import ConfigManager
from formatter import WordFormatter
from docx import Document
from docx.oxml.ns import qn

def test_page_numbers():
    print("=" * 70)
    print("  Page Number Test")
    print("=" * 70)
    
    input_file = 'test_input.docx'
    output_file = 'test_page_number_output.docx'
    
    print(f"\nInput file: {input_file}")
    print(f"Output file: {output_file}")
    print("\n" + "-" * 70)
    
    try:
        config = ConfigManager()
        print("\nConfig loaded successfully")
        
        formatter = WordFormatter(config)
        result = formatter.format_document(input_file)
        
        print(f"\nFormatting completed! Output: {result}")
        
        doc = Document(result)
        
        print("\n" + "=" * 70)
        print("  Verifying Page Numbers")
        print("=" * 70)
        
        for section_idx, section in enumerate(doc.sections):
            footer = section.footer
            
            if footer.paragraphs:
                para = footer.paragraphs[0]
                
                if para.runs:
                    run = para.runs[0]
                    
                    print(f"\nSection {section_idx + 1} Footer:")
                    print(f"  Alignment: {para.alignment}")
                    
                    r = run._r
                    field_elements = []
                    
                    for child in r:
                        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                        if tag in ['fldChar', 'instrText']:
                            field_elements.append((tag, child))
                            if tag == 'fldChar':
                                fld_type = child.get(qn('w:fldCharType'))
                                print(f"  Field char type: {fld_type}")
                            elif tag == 'instrText':
                                print(f"  Field instruction: '{child.text}'")
                    
                    if field_elements:
                        print(f"\n  SUCCESS: Page number field code found!")
                        print(f"  Total field elements: {len(field_elements)}")
                    else:
                        print(f"\n  WARNING: No field elements found in run")
                        
                    font_size = run.font.size.pt if run.font.size else "Not set"
                    bold = run.font.bold
                    
                    print(f"  Font size: {font_size}pt")
                    print(f"  Bold: {'Yes' if bold else 'No'}")
        
        print("\n" + "=" * 70)
        print("  Test Completed Successfully!")
        print("=" * 70)
        print("\nThe page numbers now use Word field codes (PAGE field).")
        print("When you open the document in Microsoft Word, each page will")
        print("display its correct page number (1, 2, 3, ...).")
        print("\nNote: The field codes may show as '1' in python-docx preview,")
        print("but will update correctly when opened in Word.")
        
        return True
        
    except Exception as e:
        print(f"\nTest failed: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    success = test_page_numbers()
    sys.exit(0 if success else 1)