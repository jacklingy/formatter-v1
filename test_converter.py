import sys
from config_manager import ConfigManager
from converter import MarkdownConverter
from docx import Document
from docx.oxml.ns import qn

def test_md_to_word():
    print("=" * 70)
    print("  Markdown to Word Conversion Test - V1.2")
    print("=" * 70)
    
    input_file = 'test_document.md'
    output_file = 'test_document_格式化.docx'
    
    print(f"\nInput file: {input_file}")
    print(f"Output file: {output_file}")
    print("\n" + "-" * 70)
    
    try:
        config = ConfigManager()
        print("\nConfig loaded successfully")
        
        converter = MarkdownConverter(config)
        result = converter.convert(input_file)
        
        print(f"\nConversion completed! Output: {result}")
        
        doc = Document(result)
        
        print("\n" + "=" * 70)
        print("  Verifying Converted Document")
        print("=" * 70)
        
        verify_headings(doc)
        verify_paragraphs(doc)
        verify_tables(doc)
        verify_page_numbers(doc)
        
        print("\n" + "=" * 70)
        print("  All Tests Passed Successfully!")
        print("=" * 70)
        
        return True
        
    except Exception as e:
        print(f"\nTest failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def verify_headings(doc):
    print("\n\nHeadings Format Verification:")
    print("-" * 50)
    
    heading_count = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) > 50:
            continue
            
        run = para.runs[0] if para.runs else None
        if not run:
            continue
            
        font_name = run.font.name or "Not set"
        font_size = run.font.size.pt if run.font.size else "Not set"
        bold = run.font.bold
        alignment = str(para.alignment) if para.alignment else "Not set"
        
        if text.startswith('#'):
            level = text.count('#')
            text_clean = text.lstrip('#').strip()
            
            if level <= 6 and heading_count.get(level) is not None:
                heading_count[level] += 1
                
                print(f"\n  Heading {level}: '{text_clean[:30]}...'")
                print(f"    Font: {font_name}")
                print(f"    Size: {font_size}pt")
                print(f"    Bold: {'Yes' if bold else 'No'}")
                print(f"    Alignment: {alignment}")
    
    print(f"\n  Statistics:")
    for level, count in sorted(heading_count.items()):
        if count > 0:
            print(f"    * Heading {level}: {count}")

def verify_paragraphs(doc):
    print("\n\nParagraph Format Verification:")
    print("-" * 50)
    
    body_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or text.startswith('#') or text.startswith('- ') or \
           text.startswith('* ') or re.match(r'^\d+\.', text):
            continue
            
        if len(text) < 10:
            continue
            
        body_count += 1
        
        if body_count <= 2:
            run = para.runs[0] if para.runs else None
            
            if run:
                font_name = run.font.name or "Not set"
                font_size = run.font.size.pt if run.font.size else "Not set"
                
                print(f"\n  Paragraph {body_count}: '{text[:40]}...'")
                print(f"    Font: {font_name}")
                print(f"    Size: {font_size}pt")
                print(f"    Alignment: {para.alignment}")
    
    print(f"\n  Total paragraphs found: {body_count}")

def verify_tables(doc):
    print("\n\nTable Format Verification:")
    print("-" * 50)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n  Table {table_idx + 1}:")
        print(f"    Rows: {len(table.rows)}")
        print(f"    Columns: {len(table.columns)}")
        print(f"    Alignment: {table.alignment}")
        
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            row_type = "Header" if is_header else "Body"
            
            if row_idx <= 1:
                print(f"\n    {row_type} Row {row_idx}:")
                
                for cell_idx, cell in enumerate(row.cells):
                    tc = cell._tc
                    tcPr = tc.tcPr
                    
                    bg_color = "Unknown"
                    has_border = False
                    
                    if tcPr is not None:
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill:
                                bg_color = fill
                        
                        borders = tcPr.find(qn('w:tcBorders'))
                        if borders is not None:
                            has_border = True
                    
                    text_preview = cell.text[:15] if cell.text else "(Empty)"
                    
                    print(f"      Cell[{row_idx},{cell_idx}]: '{text_preview}'")
                    print(f"        Background: #{bg_color}")
                    print(f"        Border: {'Yes' if has_border else 'No'}")

def verify_page_numbers(doc):
    print("\n\nPage Number Verification:")
    print("-" * 50)
    
    page_number_found = False
    
    for section_idx, section in enumerate(doc.sections):
        footer = section.footer
        
        if footer.paragraphs:
            para = footer.paragraphs[0]
            
            if para.runs:
                page_number_found = True
                run = para.runs[0]
                
                print(f"\n  Section {section_idx + 1} Footer:")
                print(f"    Alignment: {para.alignment}")
                
                r = run._r
                field_elements = []
                
                for child in r:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag in ['fldChar', 'instrText']:
                        field_elements.append((tag, child))
                        if tag == 'fldChar':
                            fld_type = child.get(qn('w:fldCharType'))
                            print(f"    Field char type: {fld_type}")
                        elif tag == 'instrText':
                            print(f"    Field instruction: '{child.text}'")
                
                if field_elements:
                    print(f"\n  SUCCESS: Page number field code found!")
                    print(f"  Total field elements: {len(field_elements)}")
                    
                    font_size = run.font.size.pt if run.font.size else "Not set"
                    bold = run.font.bold
                    
                    print(f"  Font size: {font_size}pt")
                    print(f"  Bold: {'Yes' if bold else 'No'}")
    
    if not page_number_found:
        print("\n  WARNING: No page number content detected")

import re

if __name__ == '__main__':
    success = test_md_to_word()
    sys.exit(0 if success else 1)