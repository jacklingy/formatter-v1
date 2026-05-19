from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document():
    doc = Document()
    
    doc.add_heading('文档标题测试', level=1)
    
    doc.add_paragraph('这是一段正文内容，用于测试格式化效果。')
    
    doc.add_heading('一级标题测试', level=2)
    
    doc.add_heading('二级标题测试', level=3)
    
    doc.add_heading('三级标题测试', level=4)
    
    doc.add_heading('四级标题测试', level=5)
    
    doc.add_heading('五级标题测试', level=6)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ['姓名', '年龄', '部门']
    data = [
        ['张三', '28', '技术部'],
        ['李四', '32', '市场部'],
        ['王五', '25', '财务部'],
        ['赵六', '30', '人事部']
    ]
    
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data
    
    doc.add_paragraph('\n这是第二段正文内容，包含数字 12345 和 -67890。')
    
    doc.add_page_break()
    
    doc.add_heading('第二页标题', level=1)
    
    doc.add_paragraph('这是第二页的正文内容。')
    
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'
    
    table2.rows[0].cells[0].text = '项目'
    table2.rows[0].cells[1].text = '数值'
    table2.rows[1].cells[0].text = '收入'
    table2.rows[1].cells[1].text = '12,345.67'
    table2.rows[2].cells[0].text = '支出'
    table2.rows[2].cells[1].text = '-8,901.23'
    table2.rows[3].cells[0].text = '利润'
    table2.rows[3].cells[1].text = '3,444.44'
    
    test_file = 'test_input.docx'
    doc.save(test_file)
    print(f'✅ 测试文档已创建: {test_file}')
    return test_file

if __name__ == '__main__':
    create_test_document()
