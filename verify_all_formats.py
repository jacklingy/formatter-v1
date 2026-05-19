from docx import Document
from docx.oxml.ns import qn

doc = Document('test_all_formats_格式化.docx')

print("=" * 60)
print("FORMAT VERIFICATION REPORT")
print("=" * 60)

for i, para in enumerate(doc.paragraphs[:25]):
    text = para.text[:60] if para.text else "[Empty]"
    
    formats = []
    for run in para.runs:
        run_text = run.text[:15] if run.text else ""
        if not run_text.strip():
            continue
            
        props = []
        if run.bold:
            props.append("BOLD")
        if run.italic:
            props.append("ITALIC")
        if hasattr(run.font, 'strike') and run.font.strike:
            props.append("STRIKE")
        if run.font.name == 'Consolas':
            props.append("CODE")
        
        hyperlink = run._r.getparent().getparent()
        if hyperlink.tag == qn('w:hyperlink'):
            props.append("LINK")
        
        if props:
            formats.append(f"'{run_text}'({','.join(props)})")
    
    if formats or text.strip():
        print(f"\nPara {i}: {text}")
        for fmt in formats:
            print(f"  -> {fmt}")

print("\n" + "=" * 60)
print("TABLE VERIFICATION")
print("=" * 60)

for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx + 1}: {len(table.rows)} rows x {len(table.columns)} cols")
    print(f"  Header: {[cell.text[:10] for cell in table.rows[0].cells]}")
    if len(table.rows) > 1:
        print(f"  Row 1 align: {[p.alignment for p in table.rows[1].cells[0].paragraphs]}")

print("\n" + "=" * 60)
print("VERIFICATION COMPLETE!")
print("=" * 60)
