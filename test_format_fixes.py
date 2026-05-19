# -*- coding: utf-8 -*-
"""
测试修复效果：
1. 列表项 '-' 符号残留（支持缩进列表）
2. 表格中的加粗(**)转换
3. 段落左对齐（非两端对齐）
"""

import sys
import os
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from converter import MarkdownConverter
from config_manager import ConfigManager


def test_fixes():
    print("=" * 60)
    print("测试格式转换修复")
    print("=" * 60)
    
    config = ConfigManager()
    converter = MarkdownConverter(config)
    
    test_md = """# 格式修复测试文档

## 测试1：缩进列表项（应无'-'残留，左对齐）

- 第一级项目
  - 第二级项目（缩进）
    - 第三级项目（更深缩进）
- 另一个第一级项目

## 测试2：表格中的加粗（**应转换为加粗**）

| 功能 | 状态 | 说明 |
|:-----|:----:|:-----|
| **标题层级** | ✅ | H1-H6全部支持 |
| **段落文本** | ✅ | 长短段落混合 |
| **表格功能** | ✅ | 支持对齐和样式 |

## 测试3：普通段落（应左对齐，不是两端对齐）

✅ **标题层级**：H1-H6全部使用标准Markdown语法，从一级到六级标题完整覆盖测试场景。

✅ **段落文本**：长短段落混合使用，验证不同长度文本的格式化效果和排版一致性。

✅ **表格数据**：包含多种数据类型如数字、百分比、货币等格式的正确显示。

## 测试4：正常文本段落

这是一段普通的测试文本，用于验证段落的默认对齐方式。该段落应该保持左对齐，而不是两端对齐。

这是另一段普通文本，同样应该保持左对齐的格式。
"""
    
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write(test_md)
        md_path = f.name
    
    output_path = os.path.join(tempfile.gettempdir(), 'test_format_fixes.docx')
    
    try:
        print("\n[1/2] 转换中...")
        output_path = converter.convert(md_path)
        
        if output_path:
            print(f"[2/2] 转换成功！")
            print(f"\n输出文件：{output_path}")
            
            doc = Document(output_path)
            
            print("\n" + "=" * 60)
            print("验证结果：")
            print("=" * 60)
            
            issues_found = []
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                if not text:
                    continue
                
                alignment_name = {
                    None: 'None',
                    0: 'LEFT (OK)',
                    1: 'CENTER',
                    2: 'RIGHT',
                    3: 'JUSTIFY (BAD)'
                }
                
                align_val = para.paragraph_format.alignment
                align_str = alignment_name.get(align_val, f'未知({align_val})')
                
                has_minus = text.startswith('- ') or text.startswith('• -')
                
                if '✅' in text or '**' in text or i < 15:
                    safe_text = text[:50].replace('✅', '[OK]').replace('•', '-')
                    print(f"\nPara {i+1}: {safe_text}...")
                    print(f"   对齐方式: {align_str}")
                    
                    if has_minus and not text.startswith('•'):
                        issues_found.append(f"Para{i+1}: '-' symbol remains")
                        print(f"   [WARN] List item '-' not processed!")
                    
                    if align_val == 3:
                        issues_found.append(f"Para{i+1}: Wrong justify alignment")
                        print(f"   [WARN] Should be LEFT but is JUSTIFY!")
            
            for table_idx, table in enumerate(doc.tables):
                print(f"\n表格 {table_idx + 1}:")
                for row_idx, row in enumerate(table.rows):
                    row_texts = []
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            cell_text = para.text.strip()
                            row_texts.append(cell_text[:20])
                            
                            if '**' in cell_text:
                                issues_found.append(f"Table{table_idx+1}R{row_idx+1}: Bold not converted")
                                print(f"   [WARN] '**' found in cell: {cell_text[:30]}")
                    
                    if row_texts:
                        print(f"   行{row_idx+1}: {' | '.join(row_texts)}")
            
            print("\n" + "=" * 60)
            if issues_found:
                print(f"[FAIL] Found {len(issues_found)} issues:")
                for issue in issues_found:
                    print(f"   - {issue}")
            else:
                print("[PASS] All fixes verified!")
            print("=" * 60)
            
        else:
            print("[FAIL] Conversion failed")
            
    except Exception as e:
        print(f"[ERROR] {e}")
        import traceback
        traceback.print_exc()
    finally:
        if os.path.exists(md_path):
            os.unlink(md_path)


if __name__ == '__main__':
    test_fixes()
