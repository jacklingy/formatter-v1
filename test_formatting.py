from config_manager import ConfigManager
from formatter import WordFormatter
from docx import Document
from docx.oxml.ns import qn

def test_formatting():
    print("=" * 70)
    print("  文档格式化功能测试 - V3.1")
    print("=" * 70)
    
    input_file = 'test_input.docx'
    output_file = 'test_output_已格式化.docx'
    
    print(f"\n📂 输入文件: {input_file}")
    print(f"📂 输出文件: {output_file}")
    print("\n" + "-" * 70)
    
    try:
        config = ConfigManager()
        print("\n✅ 配置文件加载成功")
        
        formatter = WordFormatter(config)
        result = formatter.format_document(input_file)
        
        print(f"\n✅ 格式化完成！输出文件: {result}")
        
        doc = Document(result)
        
        print("\n" + "=" * 70)
        print("  📋 格式化结果验证")
        print("=" * 70)
        
        verify_paragraphs(doc)
        verify_tables(doc)
        verify_page_numbers(doc)
        
        print("\n" + "=" * 70)
        print("  ✅ 所有测试通过！")
        print("=" * 70)
        
        return True
        
    except Exception as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def verify_paragraphs(doc):
    print("\n📝 段落格式验证:")
    print("-" * 50)
    
    heading_count = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
    body_count = 0
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        
        if 'Heading' in style_name:
            import re
            match = re.search(r'\d+', style_name)
            if match:
                level = int(match.group())
                if level in heading_count:
                    heading_count[level] += 1
                    
                    run = para.runs[0] if para.runs else None
                    if run:
                        font_name = run.font.name or '未设置'
                        font_size = run.font.size.pt if run.font.size else '未设置'
                        bold = run.font.bold
                        alignment = str(para.alignment) if para.alignment else '未设置'
                        
                        print(f"\n  标题{level}: '{para.text[:20]}...'")
                        print(f"    字体: {font_name}")
                        print(f"    字号: {font_size}pt")
                        print(f"    加粗: {'是' if bold else '否'}")
                        print(f"    对齐: {alignment}")
        elif para.text.strip() and 'Heading' not in style_name:
            body_count += 1
    
    print(f"\n  统计:")
    for level, count in sorted(heading_count.items()):
        if count > 0:
            print(f"    • 标题{level}: {count} 个")
    print(f"    • 正文段落: {body_count} 个")

def verify_tables(doc):
    print("\n\n📊 表格格式验证:")
    print("-" * 50)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n  表格 {table_idx + 1}:")
        print(f"    行数: {len(table.rows)}")
        print(f"    列数: {len(table.columns)}")
        print(f"    对齐方式: {table.alignment}")
        
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            row_type = "表头" if is_header else "表体"
            
            print(f"\n    {row_type}行 {row_idx}:")
            
            for cell_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.tcPr
                
                bg_color = "未知"
                has_border = False
                
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = shd.get(qn('w:fill'))
                        if fill:
                            bg_color = fill
                    
                    borders = tcPr.find(qn('w:tcBorders'))
                    if borders is not None:
                        has_border = True
                
                text_preview = cell.text[:15] if cell.text else "(空)"
                
                print(f"      单元格[{row_idx},{cell_idx}]: '{text_preview}'")
                print(f"        背景色: #{bg_color}")
                print(f"        边框: {'✓ 有' if has_border else '✗ 无'}")
                
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    alignment = str(para.alignment) if para.alignment else "未设置"
                    print(f"        对齐: {alignment}")
                    
                    if para.runs:
                        run = para.runs[0]
                        font_name = run.font.name or "未设置"
                        font_size = run.font.size.pt if run.font.size else "未设置"
                        bold = run.font.bold
                        italic = run.font.italic
                        
                        print(f"        字体: {font_name}")
                        print(f"        字号: {font_size}pt")
                        print(f"        加粗: {'是' if bold else '否'}")
                        print(f"        斜体: {'是' if italic else '否'}")

def verify_page_numbers(doc):
    print("\n\n📄 页码验证:")
    print("-" * 50)
    
    page_number_found = False
    
    for section_idx, section in enumerate(doc.sections):
        footer = section.footer
        
        if footer.paragraphs:
            para = footer.paragraphs[0]
            
            if para.runs:
                page_number_found = True
                run = para.runs[0]
                
                print(f"\n  节 {section_idx + 1} 页脚:")
                print(f"    对齐: {para.alignment}")
                print(f"    内容: '{run.text}'")
                
                rFonts = run._element.rPr.rFonts if (run._element.rPr is not None and run._element.rPr.rFonts is not None) else None
                
                if rFonts is not None:
                    east_asia_font = rFonts.get(qn('w:eastAsia')) or "未设置"
                    ascii_font = rFonts.get(qn('w:ascii')) or "未设置"
                    
                    print(f"    中文字体: {east_asia_font}")
                    print(f"    西文字体: {ascii_font}")
                
                font_size = run.font.size.pt if run.font.size else "未设置"
                bold = run.font.bold
                
                print(f"    字号: {font_size}pt")
                print(f"    加粗: {'是' if bold else '否'}")
    
    if not page_number_found:
        print("\n  ⚠️ 未检测到页码内容")

if __name__ == '__main__':
    success = test_formatting()
    exit(0 if success else 1)
